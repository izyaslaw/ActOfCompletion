from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tkinter.messagebox import *
from tkinter.filedialog import *
from dbs.dbs import *

head = \
    'АКТ о выполненной работе (по факту доставке груза)\n' \
    '№ {number}           {date}\n\n'
beginning_of_act = \
    'ИП {carrier}, именуемое в дальнейшем "Перевозчик", в лице индивидуального ' \
    'предпринимателя {carrier_gen}, действующего на основании Свидетельства о регистрации индивидуального ' \
    'предпринимателя серия {series}, и ООО "Тайди-Урал" именуемое в дальнейшем "Заказчик" в лице генерального ' \
    'директора Школык Виталия Васильевича, действующего на основании Устава с другой стороны, составили настоящий ' \
    'акт о нижеследующем:\n\n'
services = \
    '1. Заказчику были оказаны Перевозчиком следующие услуги по перевозке груза:\n\n' \
    '   Наименование услуги: Доставка товара из г.{city_from} в г.{cities_to} {dates}.\n' \
    '   Кол-во: {amount} рейс.\n' \
    '   Цена услуги: {price} рублей.\n' \
    '   ИТОГО: {price_trans} 00 копеек.\n\n' \
    '2. Услуги, определенные Поручением перевозчику, оказаны в полном объеме, качественно и в срок. ' \
    'Заказчик претензий к Перевозчику по оказанным услугам не имеет.\n\n' \
    '3. Подписи.\n\n'
signatures = 'Перевозчик (представитель)' + ' ' * 64 + 'Заказчик (представитель) '
signatures_place = '______________ {carrier}{spaces}_________________ В.В. Школык '
explanations = 'подпись, ФИО' + ' ' * 80 + 'должность, подпись, ФИО, печать'
text = ''


def createAct(main):
    data = {
        'number': main.entry_act_number.get(),
        'date': main.entry_date.get(),
        'carrier': main.pick_carrier.get(),
        'carrier_gen': getGenitive(main.carriers, main.pick_carrier.get()),
        'series': getSeries(main.carriers, main.pick_carrier.get()),
        'cities_from': createCitiesFromString(main.pick_cities_from, main),
        'cities_to': createCitiesToString(main.pick_cities_to),
        'dates': createDatesString(main.entry_dates_of_trips),
        'amount': main.entry_amount.get(),
        'price': main.entry_price.get(),
        'price_trans': createPriceTrans(main.entry_price.get())
    }
    print('2')
    document = createDocument(data)
    print('3')
    saveDocument(main, data)
    print('4')
    try:
        dir_file = open('dir.txt', 'r')
    except BaseException:
        dir_file = open('dir.txt', 'w')
        dir_file.write(askdirectory())
        dir_file.close()
        dir_file = open('dir.txt', 'r')
    dir = dir_file.read()
    dir_file.close()
    print('5')
    document_name = dir + '/АКТ {number} {carrier}.docx'.format(
        number=data['number'],
        carrier=data['carrier']
    )
    print(document_name)
    document.save(document_name)
    print('6')
    success_message = document_name + ' успешно создан!'
    showinfo(message=success_message)


def createDocument(data):
    spaces = ' ' * (80 - (len(data['carrier']) - 6) * 2)
    document = Document()

    head_paragraph = document.add_paragraph(
        head.format(
            number=data['number'],
            date=data['date']
        )
    )
    head_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    beginning_of_act_paragraph = document.add_paragraph(
        beginning_of_act.format(
            carrier=data['carrier'],
            carrier_gen=data['carrier_gen'],
            series=data['series']
        ))

    services_paragraph = document.add_paragraph(
        services.format(
            city_from=data['cities_from'],
            cities_to=data['cities_to'],
            dates=data['dates'],
            amount=data['amount'],
            price=data['price'],
            price_trans=data['price_trans']
        )
    )

    signatures_paragraph = document.add_paragraph(signatures)

    signatures_place_paragraph = document.add_paragraph(
        signatures_place.format(
            carrier=data['carrier'],
            spaces=spaces
        )
    )

    explanations_paragraph = document.add_paragraph(explanations)
    return document


def saveDocument(main, data):
    main.db_saves.insert_data(
        data['number'],
        data['date'],
        data['carrier'],
        data['cities_from'],
        data['cities_to'],
        data['dates'],
        data['amount'],
        data['price']
    )


def createCitiesFromString(cities_to, main):
    cities = [getGenitive(main.cities, city.get()) for city in cities_to]
    return ', г.'.join(cities)


def createCitiesToString(cities_to):
    cities = [city.get() for city in cities_to]
    return ', г.'.join(cities)


def getGenitive(elems, elem_name):
    for elem in elems:
        if elem_name == elem[1]:
            return elem[2]


def getSeries(carriers, carrier_name):
    for carrier in carriers:
        if carrier_name == carrier[1]:
            return carrier[3]


def createDatesString(dates):
    dates = [date.get() for date in dates]
    return ', '.join(dates)


def createPriceTrans(price_string):
    price_trans = ''
    price = int(price_string)
    thousands = price // 1000
    units = price % 1000
    if thousands > 0:
        price_trans += getWordLessThanThousand(thousands, 'тысяч ')
    price_trans += getWordLessThanThousand(units, 'рублей')
    return price_trans.capitalize()


def getWordLessThanThousand(number, word_after):
    number_trans = ''
    number_trans += getHundredsWord(number)
    if number // 10 % 10 == 1:
        number_trans += getTensIsOneWord(number)
        number_trans += word_after
    else:
        number_trans += getTensWord(number)
        if word_after == 'тысяч ':
            number_trans += getUnitsWordForThousands(number)
        else:
            number_trans += getUnitsWordForUnits(number)
    return number_trans


def getHundredsWord(number):
    hundreds = number // 100
    voc = {
        0: '',
        1: 'сто ',
        2: 'двести ',
        3: 'триста ',
        4: 'четыреста ',
        5: 'пятьсот ',
        6: 'шестьсот ',
        7: 'семьсот ',
        8: 'восемьсот ',
        9: 'девятьсот ',
    }

    return voc[hundreds]


def getTensWord(number):
    tens = number // 10 % 10
    voc = {
        0: '',
        1: '',
        2: 'двадцать ',
        3: 'тридцать ',
        4: 'сорок ',
        5: 'пятьдесят ',
        6: 'шестьдесят ',
        7: 'семьдесят ',
        8: 'восемьдесят ',
        9: 'девяносто ',
    }

    return voc[tens]


def getTensIsOneWord(number):
    units = number % 10
    voc = {
        0: 'десять ',
        1: 'одиннадцать ',
        2: 'двенадцать ',
        3: 'тринадцать ',
        4: 'четырнадцать ',
        5: 'пятнадцать ',
        6: 'шестнадцать ',
        7: 'семнадцать ',
        8: 'восемнадцать ',
        9: 'девятнадцать ',
    }

    return voc[units]


def getUnitsWordForThousands(number):
    units = number % 10

    voc = {
        0: 'тысяч ',
        1: 'одна тысяча ',
        2: 'две тысячи ',
        3: 'три тысячи ',
        4: 'четыре тысячи ',
        5: 'пять тысяч ',
        6: 'шесть тысяч ',
        7: 'семь тысяч ',
        8: 'восемь тысяч ',
        9: 'девять тысяч ',
    }

    return voc[units]


def getUnitsWordForUnits(number):
    units = number % 10

    voc = {
        0: 'рублей',
        1: 'один рубль',
        2: 'два рубля',
        3: 'три рубля',
        4: 'четыре рубля',
        5: 'пять рублей',
        6: 'шесть рублей',
        7: 'семь рублей',
        8: 'восемь рублей',
        9: 'девять рублей',
    }

    return voc[units]
